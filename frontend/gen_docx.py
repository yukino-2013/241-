"""生成角色B操作流程Word文档（按天排版·精简版）"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_title(text, level=0):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0F, 0x11, 0x15)


def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.paragraph_format.space_after = Pt(4)


def add_code(code):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)


# 封面
add_title('角色B — 前端开发日报', level=0)
add_para('适用环境：Windows + PowerShell | Python 3.10+')
add_para('')

# Day 1
add_title('Day 1：环境搭建', level=1)
add_code('mkdir E:\\smart-ops-frontend && cd E:\\smart-ops-frontend')
add_code('mkdir frontend backend docs')
add_code('python -m venv venv')
add_code('.\\venv\\Scripts\\Activate.ps1')
add_code('pip install streamlit requests')

# Day 2
add_title('Day 2：编写前端代码', level=1)
add_para('创建 frontend/app.py，实现以下功能：')
add_table(
    ['功能', '说明'],
    [
        ['聊天界面', 'Streamlit 聊天组件，支持多轮对话'],
        ['模拟模式', '后端未就绪时返回模拟数据（CPU/内存/磁盘/负载/服务）'],
        ['真实API', '调用 http://localhost:8000/chat'],
        ['快捷按钮', 'CPU / 内存 / 磁盘 / 负载 / nginx / mysql / redis'],
        ['错误处理', '连接失败、超时友好提示'],
        ['清除对话', '一键清空聊天记录'],
    ]
)

# Day 3
add_title('Day 3：启动与测试', level=1)
add_code('streamlit run frontend\\app.py')
add_para('浏览器自动打开 http://localhost:8501，验证快捷按钮和自由输入均正常返回数据。')

# Day 4
add_title('Day 4：推送到 GitHub', level=1)
add_para('⚠️ 等角色C创建好仓库后执行：')
add_code('git clone https://github.com/xxx/smart-ops-assistant.git')
add_code('cd smart-ops-assistant && git checkout -b dev')
add_code('# 复制 frontend/app.py 到项目目录')
add_code('git add frontend/app.py')
add_code('git commit -m "B: 完成Streamlit聊天界面"')
add_code('git push origin dev')

# Day 5
add_title('Day 5：与后端联调', level=1)
add_para('⚠️ 等角色A完成后端后，切换为真实模式：')
add_code('$env:USE_MOCK="false"; streamlit run frontend\\app.py')
add_para('验证前后端联调正常。')

# 日常启动
add_title('日常启动', level=1)
add_code('cd E:\\smart-ops-frontend')
add_code('.\\venv\\Scripts\\streamlit.exe run frontend\\app.py')

# 验收清单
add_title('验收清单', level=1)
add_table(
    ['#', '检查项', '对应日期', '状态'],
    [
        ['1', 'Python + 虚拟环境就绪', 'Day 1', '✅'],
        ['2', '依赖安装（streamlit, requests）', 'Day 1', '✅'],
        ['3', 'app.py 编写完成', 'Day 2', '✅'],
        ['4', 'Streamlit 服务启动，功能测试通过', 'Day 3', '✅'],
        ['5', '代码推送 GitHub dev 分支', 'Day 4', '⏳'],
        ['6', '前后端联调成功', 'Day 5', '⏳'],
    ]
)

add_para('')
add_para('文档生成时间：2026年3月31日', bold=True)

# 保存
output_path = 'D:/Desktop/角色B_前端开发日报v2.docx'
doc.save(output_path)
print(f"文档已保存到: {output_path}")
